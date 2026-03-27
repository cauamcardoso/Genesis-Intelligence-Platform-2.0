"""Genesis Intelligence Platform 2.0 — Flask Backend"""
from dotenv import load_dotenv
load_dotenv()

from flask import Flask, render_template, jsonify, request, send_file
import json
import os
import io

app = Flask(__name__)

# Load data once at startup
DATA_DIR = os.path.join(os.path.dirname(__file__), 'data')

def load_json(filename):
    path = os.path.join(DATA_DIR, filename)
    if os.path.exists(path):
        with open(path) as f:
            return json.load(f)
    return {}

faculty_data = load_json('faculty.json')
challenges_data = load_json('challenges.json')
scores_data = load_json('scores.json')
metadata = load_json('metadata.json')
advantages_data = load_json('utep_advantages.json')

# In-memory cache for AI analysis results
_analysis_cache = {}
_keywords_cache = {}

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/api/faculty")
def api_faculty():
    return jsonify(faculty_data)

@app.route("/api/challenges")
def api_challenges():
    return jsonify(challenges_data)

@app.route("/api/scores")
def api_scores():
    return jsonify(scores_data)

@app.route("/api/metadata")
def api_metadata():
    return jsonify(metadata)

@app.route("/api/advantages")
def api_advantages():
    return jsonify(advantages_data)


# ═══ Concept Generation (existing) ═══

def _format_team_for_prompt(team):
    """Shared helper to format team data for Claude prompts."""
    lines = []
    for m in team:
        line = f"- {m.get('name', 'Unknown')} ({m.get('department', '')}, {m.get('title', '').split('.')[0] if m.get('title') else 'Faculty'})"
        kws = m.get("expertise_keywords", [])[:6]
        if kws:
            line += f". Expertise: {', '.join(kws)}"
        metrics = m.get("scholar_metrics", {})
        if metrics.get("h_index"):
            line += f". h-index: {metrics['h_index']}, citations: {metrics.get('citations', 'N/A')}"
        lines.append(line)
    return "\n".join(lines)


@app.route("/api/generate-concept", methods=["POST"])
def api_generate_concept():
    try:
        import anthropic
        data = request.json
        fa = data.get("focus_area", {})
        challenge = data.get("challenge", {})
        team = data.get("team", [])
        advantages = data.get("advantages", [])

        team_text = _format_team_for_prompt(team)

        adv_text = ""
        if advantages:
            adv_text = f"\n\nUTEP Institutional Advantages:\n" + "\n".join(f"- {a}" for a in advantages)

        prompt = f"""You are helping a research university (UTEP) develop proposal concepts for the DOE Genesis Mission.

Given this focus area and research team, suggest 2-3 specific research directions that this particular team could pursue.

Challenge: {challenge.get('title', '')}
Focus Area: {fa.get('title', '')}
Focus Area Description: {fa.get('description', '')}

Research Team:
{team_text}
{adv_text}

Return ONLY valid JSON (no markdown, no code fences). Use this exact structure:
{{
  "directions": [
    {{
      "title": "Concise title under 10 words",
      "approach": "3-4 sentences describing the specific technical approach, methodology, and what it would accomplish. Be detailed and specific.",
      "team_rationale": "2 sentences on why this team's specific expertise makes them well-suited. Reference team members by name.",
      "key_questions": ["3-4 specific research questions this direction would address"],
      "potential_impact": "2 sentences on the potential scientific or practical impact if successful.",
      "suggested_next_steps": ["3-4 concrete next steps the team should take to develop this direction into a proposal"]
    }}
  ]
}}

Be specific and technical. Reference the team members' actual expertise areas. Do not use filler phrases, hedging language, or words like "leverage", "utilize", "robust", "cutting-edge", or "seamless". Write in third person. Do not use em dashes."""

        client = anthropic.Anthropic()
        message = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=2000,
            timeout=60.0,
            messages=[{"role": "user", "content": prompt}]
        )

        text = message.content[0].text.strip()
        # Try to parse as JSON
        try:
            if text.startswith("```"):
                text = text.split("```")[1]
                if text.startswith("json"):
                    text = text[4:]
            result = json.loads(text)
            return jsonify(result)
        except json.JSONDecodeError:
            # Fallback: return as plain text
            return jsonify({"concepts": text})

    except ImportError:
        return jsonify({"error": "Anthropic SDK not installed"}), 500
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(e), "type": type(e).__name__}), 500


# ═══ Opportunity Analysis (AI-powered, cached) ═══

@app.route("/api/analyze-opportunities", methods=["POST"])
def api_analyze_opportunities():
    """Analyze top focus areas and generate strategic assessments."""
    try:
        import anthropic
        data = request.json
        opportunities = data.get("opportunities", [])

        results = []
        client = anthropic.Anthropic()

        for opp in opportunities[:15]:  # Cap at 15 to control costs
            cache_key = opp.get("faId", "")
            if cache_key in _analysis_cache:
                results.append(_analysis_cache[cache_key])
                continue

            fa_title = opp.get("faTitle", "")
            fa_desc = opp.get("faDescription", "")
            challenge_title = opp.get("challengeTitle", "")
            faculty_summary = opp.get("facultySummary", "")

            prompt = f"""You are a research strategy advisor for UTEP (University of Texas at El Paso) analyzing opportunities in the DOE Genesis Mission.

Analyze this focus area opportunity and provide a strategic assessment.

Challenge: {challenge_title}
Focus Area: {fa_title}
Description: {fa_desc}

Available UTEP Faculty Pool:
{faculty_summary}

Provide exactly this JSON structure (no markdown, no code fences, just raw JSON):
{{
  "strategic_assessment": "2-3 sentences on why this is a strong or weak opportunity for UTEP. Reference specific faculty strengths.",
  "research_directions": [
    {{
      "title": "Under 10 words",
      "approach": "2-3 sentences on the technical approach.",
      "team_rationale": "1 sentence on why the available faculty are suited.",
      "alignment": "high/medium/low"
    }}
  ],
  "team_recommendation": "1-2 sentences on the ideal team composition from the available pool."
}}

Be specific. Reference faculty by name. No filler words. No "leverage", "utilize", "robust", "cutting-edge". Third person. No em dashes."""

            try:
                message = client.messages.create(
                    model="claude-sonnet-4-20250514",
                    max_tokens=1200,
                    messages=[{"role": "user", "content": prompt}]
                )
                text = message.content[0].text.strip()
                # Try to parse as JSON
                if text.startswith("```"):
                    text = text.split("```")[1]
                    if text.startswith("json"):
                        text = text[4:]
                result = json.loads(text)
                result["faId"] = cache_key
                _analysis_cache[cache_key] = result
                results.append(result)
            except (json.JSONDecodeError, Exception) as parse_err:
                results.append({
                    "faId": cache_key,
                    "strategic_assessment": f"Analysis pending. Raw response available.",
                    "research_directions": [],
                    "team_recommendation": "",
                    "_raw": str(parse_err)
                })

        return jsonify({"analyses": results})

    except ImportError:
        return jsonify({"error": "Anthropic SDK not installed"}), 500
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ═══ Literature Review Keyword Generation ═══

@app.route("/api/generate-keywords", methods=["POST"])
def api_generate_keywords():
    """Generate structured keyword sets for literature review."""
    try:
        import anthropic
        data = request.json
        fa = data.get("focus_area", {})
        challenge = data.get("challenge", {})
        team = data.get("team", [])

        cache_key = f"{fa.get('id', '')}_{'-'.join(str(m.get('faculty_id', '')) for m in team)}"
        if cache_key in _keywords_cache:
            return jsonify(_keywords_cache[cache_key])

        team_text = _format_team_for_prompt(team)

        prompt = f"""You are helping a research team prepare a literature review for a DOE Genesis Mission proposal.

Challenge: {challenge.get('title', '')}
Focus Area: {fa.get('title', '')}
Focus Area Description: {fa.get('description', '')}

Research Team:
{team_text}

Generate structured keyword sets to help this team conduct a targeted literature review. Return exactly this JSON structure (no markdown, no code fences, just raw JSON):
{{
  "core_topics": ["8-12 keywords directly related to the focus area domain"],
  "expertise_intersections": ["6-10 keywords at the overlap of team expertise and focus area needs"],
  "methodological": ["5-8 keywords for research approaches, techniques, and analytical methods"],
  "application_domains": ["5-8 keywords for end-use contexts, sectors, and application areas"],
  "search_strings": ["3-5 ready-to-paste search queries for Google Scholar or Web of Science, using Boolean operators"]
}}

Make keywords specific and technical. The expertise_intersections should reflect where THIS team's particular skills meet THIS focus area's requirements. Search strings should combine terms from multiple categories for high-precision results."""

        client = anthropic.Anthropic()
        message = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=800,
            messages=[{"role": "user", "content": prompt}]
        )

        text = message.content[0].text.strip()
        if text.startswith("```"):
            text = text.split("```")[1]
            if text.startswith("json"):
                text = text[4:]
        result = json.loads(text)
        _keywords_cache[cache_key] = result
        return jsonify(result)

    except ImportError:
        return jsonify({"error": "Anthropic SDK not installed"}), 500
    except json.JSONDecodeError:
        return jsonify({"error": "Failed to parse keyword response"}), 500
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ═══ Word Document Export ═══

@app.route("/api/export-docx", methods=["POST"])
def api_export_docx():
    """Generate a proposal package Word document."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        data = request.json
        fa = data.get("focus_area", {})
        challenge = data.get("challenge", {})
        team = data.get("team", [])
        concepts = data.get("concepts", "")
        concept_directions = data.get("concept_directions", [])
        keywords = data.get("keywords", {})
        advantages = data.get("advantages", [])
        track = data.get("track", "")
        scoring = data.get("scoring", {})

        doc = Document()

        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(10)
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # Set narrow margins
        for section in doc.sections:
            section.top_margin = Inches(0.7)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)

        # ─── Branded Header (matching AAII 2-pager style) ───
        # Title line
        title_para = doc.add_paragraph()
        run = title_para.add_run("The Institute for Applied AI Innovation (AAII)")
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x04, 0x1E, 0x42)

        # Tagline
        sub_para = doc.add_paragraph()
        run = sub_para.add_run("DOE Genesis Mission Proposal Package")
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xFF, 0x82, 0x00)
        run.font.bold = True

        # Track + FA as subtitle
        info_para = doc.add_paragraph()
        run = info_para.add_run(f"{track.replace('-', ' ').title()} Track")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run = info_para.add_run(f"  |  {fa.get('id', '')} - {fa.get('title', '')}")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # Divider
        div_para = doc.add_paragraph()
        div_para.paragraph_format.space_before = Pt(4)
        div_para.paragraph_format.space_after = Pt(8)
        run = div_para.add_run("_" * 80)
        run.font.size = Pt(6)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

        # ─── Section 1: Opportunity Context ───
        h = doc.add_heading("1. Opportunity Context", level=1)
        h.runs[0].font.color.rgb = RGBColor(0x04, 0x1E, 0x42)

        doc.add_paragraph(f"Challenge: {challenge.get('title', 'N/A')}")
        doc.add_paragraph(f"Focus Area: {fa.get('id', '')} - {fa.get('title', 'N/A')}")

        if fa.get('description'):
            doc.add_paragraph("")
            p = doc.add_paragraph(fa['description'])
            p.style.font.size = Pt(10)

        # ─── Section 2: Team Composition ───
        h = doc.add_heading("2. Proposed Team", level=1)
        h.runs[0].font.color.rgb = RGBColor(0x04, 0x1E, 0x42)

        if team:
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Light Grid Accent 1'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            headers = table.rows[0].cells
            for i, text in enumerate(["Name", "Role", "Department", "Title", "Composite"]):
                headers[i].text = text
                for p in headers[i].paragraphs:
                    for r in p.runs:
                        r.font.bold = True
                        r.font.size = Pt(10)

            for m in team:
                row = table.add_row().cells
                row[0].text = m.get("name", "")
                row[1].text = m.get("role", "").upper().replace("-", " ")
                row[2].text = m.get("department", "")
                row[3].text = m.get("title", "")
                score = m.get("composite", 0)
                row[4].text = str(round(score, 1)) if score else "N/A"
                for cell in row:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.size = Pt(10)

        # ─── Section 3: Scoring Evidence ───
        h = doc.add_heading("3. Scoring Evidence", level=1)
        h.runs[0].font.color.rgb = RGBColor(0x04, 0x1E, 0x42)

        if scoring:
            axes = [
                ("Faculty Fit", scoring.get("faculty_fit", 0), "40%"),
                ("Competitive Edge", scoring.get("competitive_edge", 0), "25%"),
                ("Team Feasibility", scoring.get("team_feasibility", 0), "20%"),
                ("Partnership Readiness", scoring.get("partnership_readiness", 0), "15%"),
            ]
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'
            headers = table.rows[0].cells
            for i, text in enumerate(["Scoring Axis", "Score (0-5)", "Weight"]):
                headers[i].text = text
                for p in headers[i].paragraphs:
                    for r in p.runs:
                        r.font.bold = True
                        r.font.size = Pt(10)
            for name, score, weight in axes:
                row = table.add_row().cells
                row[0].text = name
                row[1].text = str(round(score, 1))
                row[2].text = weight
                for cell in row:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            r.font.size = Pt(10)

            doc.add_paragraph("")
            composite = scoring.get("composite", 0)
            p = doc.add_paragraph(f"Composite Score: {round(composite, 1)} / 5.0")
            p.runs[0].font.bold = True

        # ─── Section 4: Research Directions ───
        h = doc.add_heading("4. Research Directions", level=1)
        h.runs[0].font.color.rgb = RGBColor(0x04, 0x1E, 0x42)

        if concept_directions:
            for idx, d in enumerate(concept_directions):
                # Direction title
                p = doc.add_paragraph()
                run = p.add_run(f"Direction {idx + 1}: {d.get('title', 'Untitled')}")
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x04, 0x1E, 0x42)

                # Approach
                if d.get('approach'):
                    p = doc.add_paragraph()
                    run = p.add_run("Technical Approach: ")
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run = p.add_run(d['approach'])
                    run.font.size = Pt(10)

                # Team rationale
                if d.get('team_rationale'):
                    p = doc.add_paragraph()
                    run = p.add_run("Team Positioning: ")
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run = p.add_run(d['team_rationale'])
                    run.font.size = Pt(10)

                # Key questions
                if d.get('key_questions'):
                    p = doc.add_paragraph()
                    run = p.add_run("Key Research Questions:")
                    run.font.bold = True
                    run.font.size = Pt(10)
                    for q in d['key_questions']:
                        doc.add_paragraph(q, style='List Bullet')

                # Potential impact
                if d.get('potential_impact'):
                    p = doc.add_paragraph()
                    run = p.add_run("Potential Impact: ")
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run = p.add_run(d['potential_impact'])
                    run.font.size = Pt(10)

                # Suggested next steps
                if d.get('suggested_next_steps'):
                    p = doc.add_paragraph()
                    run = p.add_run("Suggested Next Steps:")
                    run.font.bold = True
                    run.font.size = Pt(10)
                    for step in d['suggested_next_steps']:
                        doc.add_paragraph(step, style='List Bullet')

                doc.add_paragraph("")  # spacer between directions

        elif concepts:
            doc.add_paragraph(concepts)
        else:
            doc.add_paragraph("Research directions have not yet been generated for this package.")

        # ─── Section 5: Literature Review Keywords ───
        h = doc.add_heading("5. Literature Review Keywords", level=1)
        h.runs[0].font.color.rgb = RGBColor(0x04, 0x1E, 0x42)

        if keywords:
            categories = [
                ("Core Topics", keywords.get("core_topics", [])),
                ("Expertise Intersections", keywords.get("expertise_intersections", [])),
                ("Methodological", keywords.get("methodological", [])),
                ("Application Domains", keywords.get("application_domains", [])),
            ]
            for cat_name, kws in categories:
                if kws:
                    p = doc.add_paragraph()
                    run = p.add_run(f"{cat_name}: ")
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run = p.add_run(", ".join(kws))
                    run.font.size = Pt(10)

            search_strings = keywords.get("search_strings", [])
            if search_strings:
                doc.add_paragraph("")
                p = doc.add_paragraph()
                run = p.add_run("Suggested Search Strings:")
                run.font.bold = True
                for ss in search_strings:
                    doc.add_paragraph(ss, style='List Bullet')
        else:
            doc.add_paragraph("Keywords have not yet been generated for this package.")

        # ─── Section 6: UTEP Competitive Advantages ───
        if advantages:
            h = doc.add_heading("6. UTEP Competitive Advantages", level=1)
            h.runs[0].font.color.rgb = RGBColor(0x04, 0x1E, 0x42)
            for adv in advantages:
                doc.add_paragraph(adv, style='List Bullet')

        # ─── Section 7: Methodology Note ───
        h = doc.add_heading("7. Methodology Summary", level=1)
        h.runs[0].font.color.rgb = RGBColor(0x04, 0x1E, 0x42)
        doc.add_paragraph(
            "Scores were generated using multi-axis hybrid triangulation combining "
            "sentence-transformer embeddings (all-MiniLM-L6-v2) with Claude LLM evaluation. "
            "Each faculty member was scored against each focus area on four axes: "
            "Faculty Fit (40%), Competitive Edge (25%), Team Feasibility (20%), and "
            "Partnership Readiness (15%). The composite is the weighted average on a 0-5 scale. "
            "Team suggestions incorporate seniority scoring (based on academic rank, h-index, "
            "and citation percentiles) and a greedy PI allocation optimizer that maximizes "
            "coverage while respecting role constraints."
        )

        # Footer
        doc.add_paragraph("")
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("Generated by AAII x Genesis Intelligence Platform 2.0 | University of Texas at El Paso")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        fa_id = fa.get("id", "package")
        filename = f"Genesis_Proposal_{fa_id}.docx"

        return send_file(
            buffer,
            as_attachment=True,
            download_name=filename,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except ImportError:
        return jsonify({"error": "python-docx not installed"}), 500
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/health")
def health():
    return {"status": "ok"}, 200

@app.route("/api/test-claude")
def test_claude():
    """Test Claude API connectivity using raw requests first, then SDK."""
    import requests as req_lib
    results = {}

    # Step 1: Can we reach api.anthropic.com at all?
    try:
        r = req_lib.get("https://api.anthropic.com/v1/models",
                        headers={"x-api-key": os.environ.get("ANTHROPIC_API_KEY", ""),
                                 "anthropic-version": "2023-06-01"},
                        timeout=15)
        results["raw_http_status"] = r.status_code
        results["raw_http_body"] = r.text[:200]
    except Exception as e:
        results["raw_http_error"] = str(e)
        results["raw_http_type"] = type(e).__name__

    # Step 2: Try SDK
    try:
        import anthropic
        client = anthropic.Anthropic()
        message = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=50,
            messages=[{"role": "user", "content": "Say hello in one word."}]
        )
        results["sdk_ok"] = True
        results["sdk_response"] = message.content[0].text
    except Exception as e:
        results["sdk_ok"] = False
        results["sdk_error"] = str(e)
        results["sdk_type"] = type(e).__name__

    return jsonify(results)

@app.route("/api/debug-env")
def debug_env():
    key = os.environ.get("ANTHROPIC_API_KEY", "")
    has_key = bool(key)
    key_preview = key[:10] + "..." if len(key) > 10 else "(empty)"
    try:
        import anthropic
        sdk_ok = True
    except ImportError:
        sdk_ok = False
    try:
        import docx
        docx_ok = True
    except ImportError:
        docx_ok = False
    return jsonify({
        "has_api_key": has_key,
        "key_preview": key_preview,
        "anthropic_sdk": sdk_ok,
        "docx_sdk": docx_ok,
    })

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5001))
    app.run(debug=True, port=port)
